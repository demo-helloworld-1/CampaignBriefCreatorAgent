# src/tools/extract_placeholders.py

import re # Import regular expressions module
import os
from typing import List, Dict, Any # Import necessary types
from pydantic import BaseModel, Field
from langchain.tools import StructuredTool
import docx # Requires: pip install python-docx
# Import template path from config
from src.config import TEMPLATE_PATH

print("--- Defining extract_placeholders tool ---")

# --- Pydantic Schema for Tool Arguments ---
# This matches the definition in app copy.py
class ExtractPlaceholdersArgs(BaseModel):
    """Input schema for the ExtractPlaceholdersTool."""
    template_path: str = Field(description="Path to the Word template (.docx) file containing placeholders like {{PLACEHOLDER_NAME}}.")

# --- Core Python Function (Copied from app copy.py) ---
# Strictly copied the function logic as it was working
def extract_placeholders_func(template_path: str) -> Dict[str, Any]:
    """
    Extracts placeholders (like {{PLACEHOLDER_NAME}})
    from a Word document template and returns the list of unique placeholders exactly as found.

    Args:
        template_path: Path to the .docx template file.

    Returns:
        A dictionary containing:
        - 'extracted_placeholders': A list of unique placeholder strings found (exactly as in template).
        - 'status': A success or error message string.
    """
    print(f"\n--- Running extract_placeholders_func ---")
    print(f"Attempting to extract placeholders from template '{template_path}'...")
    # Use absolute path for robustness
    absolute_template_path = os.path.abspath(template_path)
    print(f"DEBUG: Absolute Template Path: {absolute_template_path}")

    if not os.path.exists(absolute_template_path):
        error_msg = f"Error: Template file not found at '{absolute_template_path}'"
        print(error_msg)
        return {"extracted_placeholders": [], "status": error_msg}

    try:
        doc = docx.Document(absolute_template_path)
        # Regex to find {{...}} allowing for whitespace inside
        placeholder_regex = re.compile(r"\{\{\s*(.*?)\s*\}\}")
        found_placeholders = set()

        # --- Function to search within text runs ---
        # This helper function finds placeholders within the runs of a paragraph or cell.
        # It reconstructs text across runs to find split placeholders.
        def find_in_runs(runs):
            full_text = "".join(run.text for run in runs)
            matches = placeholder_regex.findall(full_text)
            for match_content in matches:
                # Reconstruct the original placeholder string found
                original_placeholder = f"{{{{{match_content}}}}}"
                found_placeholders.add(original_placeholder)
                print(f"  Found '{original_placeholder}'")

        # --- Search in regular paragraphs ---
        print("Checking regular paragraphs...")
        for paragraph in doc.paragraphs:
            # Quick check to avoid processing paragraphs without {{ or }}
            if '{{' in paragraph.text and '}}' in paragraph.text:
                find_in_runs(paragraph.runs)

        # --- Search within tables ---
        print("Checking tables...")
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if '{{' in paragraph.text and '}}' in paragraph.text:
                            find_in_runs(paragraph.runs)

        if not found_placeholders:
            status_msg = f"No placeholders like {{...}} found in '{absolute_template_path}'."
            print(status_msg)
            return {"extracted_placeholders": [], "status": status_msg}
        else:
            # Convert set to list and sort for consistent output
            sorted_placeholders = sorted(list(found_placeholders))
            status_msg = f"Successfully extracted {len(sorted_placeholders)} unique placeholders from '{absolute_template_path}'."
            print(status_msg)
            print(f"Placeholders found: {sorted_placeholders}")
            return {"extracted_placeholders": sorted_placeholders, "status": status_msg}

    except Exception as e:
        error_msg = f"Error during placeholder extraction from '{absolute_template_path}': {e}"
        print(error_msg)
        import traceback
        traceback.print_exc()
        return {"extracted_placeholders": [], "status": error_msg}


# --- Create the LangChain StructuredTool (Copied from app copy.py) ---
extract_placeholders_tool = StructuredTool.from_function(
    func=extract_placeholders_func,
    name="extract_placeholders_from_template",
    description="Reads a Word document (.docx) template, extracts all unique placeholders like {{PLACEHOLDER_NAME}} found within it (in paragraphs and tables), and returns them as a list exactly as they appear in the template.",
    args_schema=ExtractPlaceholdersArgs,
    return_direct=False
)

print(f"Tool defined: {extract_placeholders_tool.name}")