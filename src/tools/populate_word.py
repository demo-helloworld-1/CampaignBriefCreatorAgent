# src/tools/populate_word.py

import os
import json # Kept import, though not used in func
from typing import Dict, Any, Optional
from pydantic import BaseModel, Field
from langchain.tools import StructuredTool
import docx
from docx.shared import Inches # Import Inches for image size
import traceback # Import traceback

# Import paths from config, though the tool accepts paths as args
# Use the paths from config for consistency, matching where the final file will be saved
from src.config import TEMPLATE_PATH, OUTPUT_PATH # Use OUTPUT_PATH from config directly

print("--- Defining populate_word tool ---")

# --- Pydantic Schema for Tool Arguments (Copied from app copy.py) ---
# This matches the definition in app copy.py
class PopulateWordArgs(BaseModel):
    """Input schema for the PopulateWordFromJSONTool."""
    json_data: Dict[str, Any] = Field(description="JSON data (Python dictionary) for text placeholders. Keys should match placeholder content (e.g., 'PLACEHOLDER_KEY' or 'KEY').")
    image_placeholders: Optional[Dict[str, str]] = Field(default=None, description="Optional dictionary mapping image placeholder content (e.g., 'PLACEHOLDER_COMPANY_LOGO') to the local image file path (e.g., './logos/nike.png').")
    template_path: str = Field(description="Path to the Word template (.docx) file containing placeholders.")
    output_path: str = Field(description="Path where the populated Word document will be saved.")


# --- Core Python Function (Copied from app copy.py) ---
# Strictly copied the function logic as it was working
def populate_word_from_json_func(
    json_data: Dict[str, Any],
    template_path: str,
    output_path: str,
    image_placeholders: Optional[Dict[str, str]] = None
    ) -> str:

    """
    Populates a Word document template with text data from json_data
    and image data from image_placeholders.

    Text Placeholders: {{PLACEHOLDER_KEY}} or {{KEY}}. Keys in json_data match 'PLACEHOLDER_KEY' or 'KEY'.
    Image Placeholders: {{PLACEHOLDER_IMAGE_KEY}}. Keys in image_placeholders match 'PLACEHOLDER_IMAGE_KEY'.

    Args:
        json_data: Dictionary containing the text data.
        template_path: Path to the .docx template file.
        output_path: Path to save the populated .docx file.
        image_placeholders: Dictionary mapping image placeholder content to image file paths.

    Returns:
        A success or error message string.
    """
    print(f"\n--- Running populate_word_from_json_func ---")
    print(f"Attempting to populate template '{template_path}' and save to '{output_path}'...")

    # --- Input Validation ---
    if not isinstance(json_data, dict):
        msg = f"ERROR: Input 'json_data' is not a dictionary (received type: {type(json_data)})."
        print(msg)
        return msg
    # Validate type only if not None
    if image_placeholders is not None and not isinstance(image_placeholders, dict):
        msg = f"ERROR: Input 'image_placeholders' must be a dictionary or None (received type: {type(image_placeholders)})."
        print(msg)
        return msg

    image_placeholders = image_placeholders or {} # Default to empty dict if None

    print(f"DEBUG: Received json_data keys: {list(json_data.keys())}")
    print(f"DEBUG: Received image_placeholders keys: {list(image_placeholders.keys())}")

    # --- Path handling ---
    absolute_template_path = os.path.abspath(template_path)
    absolute_output_path = os.path.abspath(output_path)
    print(f"DEBUG: Absolute Template Path: {absolute_template_path}")
    print(f"DEBUG: Absolute Output Path: {absolute_output_path}")

    try:
        if not os.path.exists(absolute_template_path):
            return f"Error: Template file not found at '{absolute_template_path}'"

        doc = docx.Document(absolute_template_path)
        text_keys_successfully_replaced = set()
        image_keys_successfully_replaced = set()

        # --- Helper Function to find and replace TEXT ---
        # This function iterates through the runs of a paragraph/cell
        # and replaces placeholder text. It handles cases where placeholders
        # might be split across runs due to formatting.
        def find_and_replace_text_in_runs(paragraph_or_cell_runs, data_dict):
            """Finds and replaces text placeholders within a list of runs."""
            nonlocal text_keys_successfully_replaced
            # Get combined text of the runs to search for placeholders
            combined_text = "".join(r.text for r in paragraph_or_cell_runs)
            replaced_text = combined_text # Start with the combined text

            json_keys = list(data_dict.keys())

            for key_from_json in json_keys:
                # Construct the two possible placeholder formats
                # Ensure the key is treated correctly if it already starts with PLACEHOLDER_
                content_key = key_from_json # The key from the JSON data (e.g., "CAMPAIGN_NAME" or "PLACEHOLDER_COMPANY_LOGO")

                # Construct placeholder to find: {{KEY}} or {{PLACEHOLDER_KEY}}
                # The original code's logic for constructing placeholder_to_find was slightly complex
                # Let's replicate the original logic as closely as possible based on the app copy.py snippet
                if not content_key.upper().startswith("PLACEHOLDER_"):
                    placeholder_to_find_1 = f"{{{{{content_key}}}}}" # e.g., {{KEY}}
                    placeholder_to_find_2 = f"{{{{PLACEHOLDER_{content_key}}}}}" # e.g., {{PLACEHOLDER_KEY}}
                else:
                    placeholder_to_find_1 = f"{{{{{content_key}}}}}" # e.g., {{PLACEHOLDER_KEY}}
                    placeholder_to_find_2 = None # Only look for {{PLACEHOLDER_...}} if key starts with it


                text_to_insert = str(data_dict.get(key_from_json, "")) # Get the value to insert, default to empty string

                # Check if either placeholder exists in the combined text
                placeholder_found = False
                placeholder_to_use = None # Keep track of which placeholder was found/used

                if placeholder_to_find_1 and placeholder_to_find_1 in replaced_text:
                    placeholder_to_use = placeholder_to_find_1
                    placeholder_found = True
                elif placeholder_to_find_2 and placeholder_to_find_2 in replaced_text:
                    placeholder_to_use = placeholder_to_find_2
                    placeholder_found = True
                else:
                    continue # Neither placeholder found, skip this key


                if placeholder_found and placeholder_to_use: # Ensure placeholder_to_use is not None
                    print(f"  Attempting to replace TEXT '{placeholder_to_use}' with '{text_to_insert[:50]}...' (using key '{key_from_json}')")

                    # Find the location(s) of the placeholder in the combined text
                    start_index = replaced_text.find(placeholder_to_use)
                    while start_index != -1:
                        end_index = start_index + len(placeholder_to_use)

                        # Replace in the combined text string
                        replaced_text = replaced_text[:start_index] + text_to_insert + replaced_text[end_index:]

                        # Find the next occurrence - start search after the newly inserted text
                        start_index = replaced_text.find(placeholder_to_use, start_index + len(text_to_insert))


                    # Now, update the runs based on the final replaced_text.
                    # This simplifies run structure but replaces text reliably.
                    if paragraph_or_cell_runs:
                        paragraph_or_cell_runs[0].text = replaced_text # Put all text in the first run
                        # Clear the text from all other runs
                        for run in paragraph_or_cell_runs[1:]:
                            run.text = ""

                    print(f"  Successfully replaced TEXT for '{placeholder_to_use}' using key '{key_from_json}'")
                    text_keys_successfully_replaced.add(key_from_json)
                    # After replacing for this key, the text is updated for subsequent key searches within the same runs


        # --- Helper Function to find and replace IMAGES (Copied from app copy.py) ---
        # Strictly copied the function logic
        def find_and_replace_image_in_runs(paragraph_or_cell_runs, image_data_dict):
            """Finds image placeholders and inserts images."""
            nonlocal image_keys_successfully_replaced
            # Iterate through image placeholders provided
            # image_data_dict keys are the content *inside* the placeholder, e.g., 'PLACEHOLDER_COMPANY_LOGO'
            for img_placeholder_content, img_path in image_data_dict.items():
                # Construct the full placeholder string to search for
                placeholder_to_find = f"{{{{{img_placeholder_content}}}}}" # e.g., "{{PLACEHOLDER_COMPANY_LOGO}}"

                # Check if placeholder exists in the text combined from runs
                combined_text = "".join(r.text for r in paragraph_or_cell_runs)

                if placeholder_to_find in combined_text:
                    print(f"  Found IMAGE placeholder '{placeholder_to_find}'. Attempting insertion...")
                    # Verify the image file exists
                    if not os.path.exists(img_path):
                        print(f"  WARNING: Image file not found at path: '{img_path}' for placeholder '{placeholder_to_find}'. Skipping.")
                        continue # Skip to next image key

                    # Find the run(s) containing the placeholder.
                    # Simple approach: Find the first run, clear its text, insert picture.
                    found_run_index = -1
                    for i, run in enumerate(paragraph_or_cell_runs):
                        if placeholder_to_find in run.text:
                            found_run_index = i
                            break # Found the first run containing the placeholder

                    if found_run_index != -1:
                        run_to_replace = paragraph_or_cell_runs[found_run_index]
                        original_run_text = run_to_replace.text # Keep original text to replace only the placeholder

                        try:
                            # Replace the placeholder text within the run's text
                            run_to_replace.text = original_run_text.replace(placeholder_to_find, '')
                            print(f"  Cleared placeholder text in run {found_run_index}.")

                            # Add the picture to the run
                            # Use Inches(1.5) as in original code
                            run_to_replace.add_picture(img_path, width=Inches(1.5))
                            print(f"  Inserted IMAGE '{img_path}' into run {found_run_index}.")
                            image_keys_successfully_replaced.add(img_placeholder_content)

                        except Exception as img_e:
                            print(f"  ERROR: Failed to insert image '{img_path}' for placeholder '{placeholder_to_find}': {img_e}")
                            traceback.print_exc() # Print image insertion error traceback
                    else:
                        # This could happen if the placeholder spans multiple runs
                        print(f"  WARNING: Could not find a single run containing image placeholder '{placeholder_to_find}'. Placeholder might span runs or be complex.")

        # --- Main Processing Loop (Copied from app copy.py) ---
        print("Checking paragraphs and tables for text and images...")
        all_content_items = [] # Collect paragraphs and table cells

        # Collect all paragraphs
        all_content_items.extend(doc.paragraphs)

        # Collect all paragraphs within table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_content_items.extend(cell.paragraphs) # Each cell can contain multiple paragraphs

        # Process each paragraph/cell content
        for item in all_content_items:
            # Process text replacements first
            find_and_replace_text_in_runs(item.runs, json_data)
            # Then process image replacements if image data is provided
            if image_placeholders:
                find_and_replace_image_in_runs(item.runs, image_placeholders)


        # --- Reporting (Copied from app copy.py) ---
        print("\n--- Replacement Summary ---")
        # Text reporting
        provided_text_keys = set(json_data.keys())
        text_keys_not_found_in_template = provided_text_keys - text_keys_successfully_replaced
        text_placeholders_replaced_count = len(text_keys_successfully_replaced)

        print(f"Text placeholders replaced: {text_placeholders_replaced_count} out of {len(provided_text_keys)} provided text keys.")
        if text_keys_not_found_in_template:
            print(f"WARNING: Text values provided but placeholders NOT found in template for keys: {sorted(list(text_keys_not_found_in_template))}")

        # Image reporting
        provided_image_keys = set(image_placeholders.keys())
        image_keys_not_found_in_template = provided_image_keys - image_keys_successfully_replaced
        image_placeholders_replaced_count = len(image_keys_successfully_replaced)

        print(f"Image placeholders replaced: {image_placeholders_replaced_count} out of {len(provided_image_keys)} provided image keys.")
        if image_keys_not_found_in_template:
            print(f"WARNING: Image paths provided but placeholders NOT found in template for keys: {sorted(list(image_keys_not_found_in_template))}")


        # --- Directory creation and saving logic (Copied from app copy.py) ---
        # Get output directory from the calculated absolute path
        output_dir = os.path.dirname(absolute_output_path)
        # Use exist_ok=True for robustness
        if output_dir and not os.path.exists(output_dir):
            try:
                os.makedirs(output_dir, exist_ok=True)
                print(f"Created output directory: {output_dir}")
            except OSError as dir_e:
                msg = f"Error creating output directory '{output_dir}': {dir_e}"
                print(msg)
                return msg # Stop if directory creation fails

        print(f"\nAttempting to save populated document to: '{absolute_output_path}'...")
        try:
            doc.save(absolute_output_path)
            print(f"Successfully called doc.save() for: {absolute_output_path}")
            # Optional: Add a check if the file actually exists after saving
            if not os.path.exists(absolute_output_path):
                msg = f"Error: File saving failed silently for {absolute_output_path}"
                print(msg)
                return msg
            else:
                print(f"File confirmed to exist at '{absolute_output_path}' after saving.")
                # Return the success message as in app copy.py
                return f"Successfully populated template (Text & Images) and saved to '{absolute_output_path}'"
        except Exception as e_save:
            msg = f"Error during file save operation to '{absolute_output_path}': {e_save}"
            print(msg)
            traceback.print_exc() # Print save error traceback
            return msg

    except FileNotFoundError:
        msg = f"Error: Template file not found at '{absolute_template_path}' during processing."
        print(msg)
        return msg
    except Exception as e:
        msg = f"An unexpected error occurred during Word population from '{absolute_template_path}': {e}"
        print(msg)
        traceback.print_exc() # Print general processing error traceback
        return msg


# --- Create the LangChain StructuredTool (Copied from app copy.py) ---
populate_word_tool = StructuredTool.from_function(
    func=populate_word_from_json_func,
    name="populate_word_from_json",
    description="Populates a Word (.docx) template with text and images. Requires 'json_data' (dict for text placeholders like {{PLACEHOLDER_KEY}}), 'template_path', 'output_path', and optionally 'image_placeholders' (dict mapping placeholder content like 'PLACEHOLDER_COMPANY_LOGO' to image file paths). Text keys in json_data match content inside braces (e.g., 'PLACEHOLDER_KEY' or 'KEY'). Image keys in image_placeholders match image placeholder content (e.g., 'PLACEHOLDER_COMPANY_LOGO').",
    args_schema=PopulateWordArgs,
    return_direct=False
)

print(f"Tool defined: {populate_word_tool.name}")